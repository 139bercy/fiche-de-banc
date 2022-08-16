

#maths
import networkx as nx
import numpy as np
import math
import matplotlib.pyplot as plt
import pandas as pd
from random import randrange
#import scipy

#NLP
from convectors.layers import Tokenize, Snowball, OneHot, TfIdf, SVD, Lemmatizer, CountVectorizer
from scipy.spatial.distance import pdist, squareform

import gensim
from gensim.utils import simple_preprocess
from gensim.parsing.preprocessing import STOPWORDS
from gensim import corpora, models, similarities
from gensim import matutils
from gensim.similarities.docsim import MatrixSimilarity
from gensim.corpora import Dictionary
from gensim.models import TfidfModel

from nltk.stem import WordNetLemmatizer, SnowballStemmer
from nltk.stem.porter import *
from nltk.corpus import stopwords

from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.feature_extraction.text import CountVectorizer

from wordcloud import WordCloud
#from PIL import Image

#ouvrir fichiers et pre processing
import fitz #pdf
from os import listdir
from os.path import isfile, join
import html2text
#import string
#import unidecode

#creer docx
import os
from docx import Document
from docx.enum.text import WD_COLOR_INDEX
import shutil

#gestion param dans fonctions
from typing import Literal, Union


#a trier
from mrakun import RakunDetector
from time import time
from tqdm import tqdm

from time import mktime, strptime
from datetime import datetime

import nltk
from nltk.corpus import stopwords
from nltk.cluster.util import cosine_distance
import numpy as np
import pandas as pd
import networkx as nx
import spacy
import fr_core_news_sm
import fr_core_news_md
import fr_core_news_lg
from sklearn.metrics.pairwise import cosine_similarity
import networkx as nx
from nltk.stem.snowball import SnowballStemmer

import torch
from scipy.spatial.distance import cosine
from sklearn.manifold import TSNE
from sklearn.cluster import KMeans
import matplotlib.pyplot as plt
import plotly.express as px
import fr_core_news_md


from sentence_transformers import SentenceTransformer

import importlib
#import logger
from wordcloud import WordCloud
from docx import Document
from docx.shared import Inches

#install
#python -m spacy download fr_core_news_lg


# In[2]:


#!python -m spacy validate


# # 1 Import des textes <a name="import"></a>
# 
# [Retour table des matières](#up)

# In[3]:


def text_from_pdf(path : str):
    """
    Extrait le texte et si possible la date d'un document pdf
    
    paramètres
    -----
    path : str
        chemin de la localisation du document
    
    returns
    -----
    text : str
        texte du document
    dt : datetime.datetime ou None
        date du doc sous format datetime ou None si la date n'a pas pu être récupérée.
    
    count : int
        nombre de pages du document
    """
    with fitz.open(path) as doc:
        try:
            datestring = doc.metadata['creationDate'][2:-7] #récupérer la date
            ts = strptime(datestring, "%Y%m%d%H%M%S")
            dt = datetime.fromtimestamp(mktime(ts)) #convertir au format datetime
        except:
            dt = None
        text = ""
        count = 0
        for page in doc:
            text += page.get_text(flags=None)
            count +=1
    return text, dt, count


# In[4]:


#possible de récup la date ?
def text_from_html(path : str):
    """
    Extrait le texte d'un document html
    
    paramètres
    -----
    path : str
        chemin de la localisation du document
    
    returns
    -----
    text : str
        texte du document
    dt : None
        renvoie toujours none pour la date 
    """
    with open(path, "r") as doc:
        text = doc.read()
        res = html2text.html2text(text)
        dt = None
    return res, dt


# In[5]:


def import_data(directory, verbose=True):
    """
    Création du corpus
    On ne garde que les textes de plus de 1500 caractères (en dessous c'est souvent des tableaux ou des plannings)
    et moins de 6 pages (au dessus ce ne sont pas des tracts)
    
    paramètres
    -----
    directory : str
        chemin du dossier contenant tous les documents à importer
    verbose : bool
        si True, affiche des informations supplémentaire
        
    returns
    -----
    titres, raw_text, datetime, type_doc : list
        liste des titres(str), texte brut(str), dates(datetime) et type(str) des documents initiaux
    """
    fichiers = [f for f in listdir(directory) if isfile(join(directory, f))]
    if verbose:
        print(f"{len(fichiers)} documents")
    titres = []
    raw_text = []
    datetime = []
    type_doc = []
    suppr = 0
    for f in fichiers:
    #IMPORT DES PDF
        if f.endswith('.pdf'):
            text, dt, count = text_from_pdf(directory + f)
            if len(text) <= 1500: #ne garde pas les docs trop courts
                if verbose:
                    print(f'texte pdf supprimé (trop court {len(text)} caractères):{f}') 
                suppr += 1
            elif count >= 6 :
                if verbose:
                    print(f'texte pdf supprimé (trop long {count} pages):{f}') 
                suppr += 1
            else: 
                raw_text.append(text)
                titres.append(f)
                datetime.append(dt)
                type_doc.append('pdf')
        elif f.endswith('.html'):
    #IMPORT DES HTML
            text, dt = text_from_html(directory + f)
            if len(text) <= 1500: #ne garde pas les docs trop courts
                if verbose:
                    print(f'texte html supprimé:{f}')
                suppr += 1
            else: 
                raw_text.append(text)
                titres.append(f)
                datetime.append(dt)
                type_doc.append('html')
        else:
            if verbose:
                print(f"INCONNU:{f}")
            suppr += 1
    if verbose:        
        print(f"au total, {suppr} textes supprimés")
        print(f"{len(raw_text)} documents traités")
    return titres, raw_text, datetime, type_doc


# In[6]:


#Si besoin de trouver l'indice d'un doc
#A SUPPRIMER PLUS TARD
def find_doc_id(string):
    for i in range(len(titres)):
        if titres[i].startswith(string):
            print(i)
            print(titres[i])
            break


# In[7]:


def text_preproc(x):
    """
    pre-processing d'un text
    
    tâches éffectuées : 
        retirer les coupures lorsqu'un mot apparaissait en fin de ligne dans le document original;
        retirer les liens (qui commencent par http) : souvent des images
        retirer les sauts de lignes simples (qui correspondent à un retour à la ligne dans un pdf)
        retirer les images
        ajouter un espace après les sauts de ligne si un mot apparait directement après
        
    paramètre
    -----
    x : str
        texte brut
    
    returns
    -----
    x : str
        texte traité
    """
    #x = x.lower()
    x = re.sub(r'-\n', '', x) #enlever les coupures au milieu des mots
    #suppression des liens (même lorsqu'il y a \n au milieu !)
    x = re.sub(r'\n', 'NEWLINE', x)
    x = re.sub(r'http\S+', '', x) #retirer les liens:\S s'arrete s'il y a un retour à la ligne donc on le remplace avant
    x = re.sub(r'NEWLINE', '\n', x)
    #retirer retours à la ligne
    x = re.sub(r'\n\n', 'NEWLINE', x) #on remplace les vrais retours à la ligne
    x = re.sub(r'\n', ' ', x) #on supprime les faux, ie fin de ligne du pdf
    x = re.sub(r'NEWLINE', '\n\n', x) #on remet les vrais
    #x = x.encode('ascii', 'ignore').decode()
    #x = re.sub(r'@\S+', '', x)
    #x = re.sub(r'#\S+', '', x)
    #x = re.sub(r'\'', '', x)
    x = re.sub(r'\S+image\S+', '', x)
    x = re.sub(r'\n(?=\S+)', ' ', x)
    return x


# # 2 Initialisation des tokenizer <a name="init"></a>
# 
# [Retour table des matières](#up)

# In[8]:


def init_tokenizer():
    """
    initialise les tokenizer pour être certains d'utiliser les mêmes tout le long
    
    returns 
    -----
    sentence_tok : tokenize un document à l'echelle des phrases
    
    word_tok : tokenize un document à l'echelle des mots
    """
    sentence_tok = Tokenize(sentence_tokenize=True,
                                  word_tokenize=False,
                                  strip_accents=False,
                                  lower=False,
                                  verbose=False)

    word_tok = Tokenize(stopwords=["fr"], word_tokenize=True, strip_accents=False, lower=True, verbose=False)
    word_tok += Lemmatizer()
    #word_tok += Snowball(lang="french") #stemming
    return sentence_tok, word_tok


# # 3 Construction du dictionnaire sur tout le corpus <a name="dico"></a>
# 
# [Retour table des matières](#up)

# In[9]:


def build_dict_sentence(corpus, sentence_tok, word_tok):
    '''
    Construction d'un dictionnaire corpora sur tout le corpus.
    Il contiendra notamment les fréquences par phrases.
    Cela servira à calculer le idf -inverse document frequency- sur tout le corpus.
    
    param
    ------
    corpus:list
        le corpus en entier: liste de plusieurs documents de plusieurs phrases
    sentence_tok:function (convectors.model)
    word_tok:function (convectors.model)
    
    returns
    -----
    dictionary:gensim.corpora.Dictionary
        dictionnaire de tous les mots de tout le corpus contenant notamment les fréquences par phrase
    
    '''
    #corpus_sentences est une série contenant la liste des phrases par doc
    corpus_sentences = sentence_tok(corpus)
    
    #construction de la liste des phrases tout doc confondus
    bag_of_sentences = []
    for doc in corpus_sentences:
        for sentence in doc:
            bag_of_sentences.append(sentence)
            
            
    #taille moyenne d'une phrase (en nb de caractères)
    avg = sum(map(len, bag_of_sentences))/float(len(bag_of_sentences))
    #print(f"les phrases font en moyenne {avg} caractères.")
    
    #words est une série contenant la liste des mots par phrase
    words = word_tok(bag_of_sentences)
    
    dictionary = gensim.corpora.Dictionary(words)
    
    return dictionary


# In[10]:


def build_dict_document(corpus, word_tok):
    """
    dictionnaire de tous les mots mais avec les fréquences par document.
    Contient les mêmes mots que build_dict_sentences mais pas les memes freq
    """
    
    words_per_doc = word_tok(corpus) #liste de liste de mots words_per_doc[i] = liste des mots du doc 
    dictionary = gensim.corpora.Dictionary(words_per_doc)
    
    return dictionary


# In[11]:


def highest_frequency(dictionary, n=None):
        """
        renvoie la liste des mots les plus frequents
        
        paramètres
        ----------
        n:int or None, optional
            The number of most common words to be returned. If `None`, all words in the dictionary
            will be returned
        Returns
        -------
        most_common:list of (str, int)
            The n most common words and their counts from the most common to the least.
        """
        most_common = [
            (dictionary[word], count)
            for word, count
            in sorted(dictionary.cfs.items(), key=lambda x: (-x[1], x[0]))[:n]
        ]
        return most_common


# In[12]:


def lowest_frequency(dictionary, n= None):
        """
        renvoie la liste des mots les moins frequents
        
        paramètres
        ----------
        n:int or None, optional
            The number of most common words to be returned. If `None`, all words in the dictionary
            will be returned. 
        Returns
        -------
        less_common:list of (str, int)
            The n most common words and their counts from the most common to the least.
        """
        less_common = [
            (dictionary[word], count)
            for word, count
            in sorted(dictionary.cfs.items(), key=lambda x: (-x[1], x[0]), reverse=True)[:n]
        ]
        return less_common


# In[13]:


def most_present(dictionary, n= None):
    """
    renvoie la liste des mots qui apparaissent dans le plus de documents
    """
    most_common = [
        (dictionary[word], count)
        for word, count
        in sorted(dictionary.dfs.items(), key=lambda x: (-x[1], x[0]))[:n]
    ]
    return most_common


# In[14]:


#highest_frequency(dictionary, n= None)


# In[15]:


def less_present(dictionary, n= None):
    """
    renvoie la liste des mots qui apparaissent dans le plus de documents
    """
    less_common = [
        (dictionary[word], count)
        for word, count
        in sorted(dictionary.dfs.items(), key=lambda x: (-x[1], x[0]), reverse=True)[:n]
    ]
    return less_common


# # 4 Construction de la matrice TFIDF <a name="tfidf"></a>
# 
# [Retour table des matières](#up)

# In[16]:


def compute_tfidf(sentences, word_tok, method: Union[Literal["local"], Literal["global"]], dictionary=None):
    '''
    Calcul de la matrice tfidf pour chacune des phrases d'un doc
    Devra utiliser le dictionnaire fait sur tous les docs(si global)
    
    paramètres:
    sentences: liste des phrases de 1 document
    dictionary :corpora.dictionary
    method: local calcule le tfidf par rapport aux phrases, et global par rapport aux documents
    '''
    if method == "global":
        if not dictionary :
            print("Il faut ajouter un dictionnaire sur les documents")
        #word_list est la liste des mots par phrase
        word_list = word_tok(sentences)
        bow = []
        for sentence in word_list:
            bow.append(dictionary.doc2bow(sentence))
        tfidf = models.TfidfModel(dictionary=dictionary) #fit dictionary
        result = [] #quel type choisir ?
        for sent in bow: #pour chaque phrase
            result.append(tfidf[sent])
        mat = matutils.corpus2csc(result,num_docs=len(result)) #convert to sparse array  
        X = mat.T #transposée 
    else:
        if dictionary:
            print("vous avez renseigné un dictionnaire alors que vous utilisez une méthode locale qui n'en nécessite pas")
        nlp = word_tok.copy()
        nlp += TfIdf()
        nlp.verbose = False
    
        X = nlp(sentences)
           
    return X


# In[17]:


#corpus = [" mot0 mot1 mot2 mot3. mot1 mot4 mot5 mot3. mot2 mot6 mot1 mot1 mot4 mot5 mot3.",
#         "mot2 mot6 mot7 mot2. mot1. mot1 mot2. mot1 mot4 mot5 mot3. mot7 mot8 mot9."]


# # 5 Text Rank  <a name="tr"></a>
# 
# [Retour table des matières](#up)

# In[18]:


def summarize_tr(text, dictionary = None, n_sentences=5, ratio=0, method: Union[Literal["local"], Literal["global"]] = 'local', nlp=None):   
    '''
    Renvoie les phrases les plus importantes d'un doc en utilsant textrank (algo de graphes)
    si method = local
    Compare les phrases d'un document en interne (tfidf calculé par rapport à un doc seulement)
    si method = global
    Compare les phrases par rapport à tous les doc du corpus (sur lequel le dictionnaire a été construit)

    param: 
    ------
    text:string
        un texte du corpus
    dictionary:
        dictionnaire sur tout le corpus
    n_sentences:int
        nombre de phrases max souhaitées pour le résumé
    ratio:float entre 0 et 1
        proportion de phrases souhaitée pour le résumé
    method:'global' ou 'local'
        si method = local
        Compare les phrases d'un document en interne (tfidf calculé par rapport à un doc seulement)
        si method = global
        Compare les phrases par rapport à tous les doc du corpus (sur lequel le dictionnaire a été construit)
        
    nlp : str
        utilisation d'une base pré entrainée et PLUS de textrank
        
    returns
    ------
    sentences:
        liste des phrases (tokenisées) du document initial
    summary:
        liste des phrases du résumé
    sentence_id
        liste des indices des phrases du résumé (pour pouvoir récupérer les originales)

    '''
    sentence_tok, word_tok = init_tokenizer()

    if (method == "global") and (dictionary == None):
        print("Vous devez ajouter un dictionnaire pour la méthode globale")
        return
    
    sentences = sentence_tok([text])
    sentences = sentences[0]
    for s in sentences:
        if s == '':
            sentences.remove(s)
    
    #Construction de la matrice de similarité
    if nlp == None:
        X = compute_tfidf(sentences, word_tok, method, dictionary)
        sim = X @ X.T #cosine similarity, pas besoin de normaliser car tfidf
        G = nx.from_scipy_sparse_matrix(sim)
    else:
        sim = similarity_matrix(sentences, nlp)
        G = nx.from_numpy_array(sim)


    pagerank = nx.pagerank(G, max_iter=100)
    # chaque phrase obtient un score
    pagerank = sorted(pagerank.items(), key=lambda x: x[1], reverse=True)
    # pagerank contient des couples (phrase, score)
    
    if ratio > 0 and ratio < 1: #utilisation d'un ratio plutot qu'un nb de phrases à récup
        n_sentences = math.floor(ratio * len(sentences))

    summary = []
    sentence_id = []
    for index, _ in pagerank:
        sentence = sentences[index]
        if len(sentence) < 50:  # avoid too short sentences
            continue
        sentence_id.append(index)
        sentence = sentence.strip().replace(" \n", " ").replace("\n", " ") #VERIFIER ?
        summary.append(sentence.strip())

        if len(summary) == n_sentences:
            break
          
    #print(f"taille du résumé: {len(summary)}")
    return sentences, summary, sentence_id


# In[19]:


#sentences, summary, sentence_id = summarize_tr(data[0], dictionary = dictionary, n_sentences=5, ratio=0, method='global')    


# # 6 Affichage docx <a name="docx"></a>
# 
# [Retour table des matières](#up)

# In[20]:


def create_save_file(doc, title, type_file, comm='', folder=None, move_original=False):
    """
    Sauvegarde le document dans un nouveau (s'il n'existe pas déjà) dossier du même nom.
    Rajoute eventuellement dans ce même dossier le pdf original.
    
    paramètres
    -----
    doc : docx.Document
        document à sauvegarder  
    title : str
        titre du document
    type_file : str
        type du document
    comm : str
        commentaire à ajouter au titre
    folder : str
        titre du dossier si on le veut différent du titre
    move_original : bool
        si true, ajoute le pdf original dans le dossier
    
    """
    if folder==None:
        folder = title
        
    #creation dossier output
    if not(os.path.exists('output/')):
        os.makedirs('output/')
        
     #creation dossier title
    directory = f'output/{folder}'
    if not(os.path.exists(directory)): #si le dossier n'existe pas déjà, on le créé
        os.makedirs(directory)
        
    name = (f'{directory}/{title}_{comm}.{type_file}')
    
    if move_original:
        if not(os.path.exists(directory+title)):
            try: #ajout du vrai pdf dans le dossier output pour comparer
                shutil.copy('archives/'+title, directory)
            except:
                print("Le pdf n'a pas été trouvé")        
    
    doc.save(name)


# In[21]:


def summarize_docx(text, titre = None, comm=None, dictionary = None, n_sentences=5, ratio=0, nlp=None, method: Union[Literal["local"], Literal["global"]] = 'local'):
    '''
    créé un résumé + le sauvegarde dans un docx
    
    créé un docx contenant toutes les phrases du document d'origine (sans le layput avec les images)
    et avec les phrases importantes choisies par l'algo surlignées.
    Enregistre le docx dans un nouveau dossier du même nom
    
    paramètres
    -----
    text:string
        un texte du corpus, le texte à résumer
    titre : str
        titre à donner au docx sauvegardé
    comm : str
        commentaire à ajouter au titre
    dictionary:
        dictionnaire sur tout le corpus
    n_sentences:int
        nombre de phrases max souhaitées pour le résumé
    ratio:float entre 0 et 1
        proportion de phrases souhaitée pour le résumé
    method:'global' ou 'local'
        si method = local
        Compare les phrases d'un document en interne (tfidf calculé par rapport à un doc seulement)
        si method = global
        Compare les phrases par rapport à tous les doc du corpus (sur lequel le dictionnaire a été construit)
        
    nlp : str
        utilisation d'une base pré entrainée et PLUS de textrank
    
    returns
    -----
    summary : résumé créé par summarize_tr
    '''
    
    sentences, summary, sentences_id = summarize_tr(text, dictionary, n_sentences, ratio, method, nlp=nlp)

    doc_result = Document()
    paragraph = doc_result.add_paragraph()
    for i in range(len(sentences)):
        sentence = sentences[i]
        run = paragraph.add_run(sentence.strip())
        if i in sentences_id:
            run.bold = True
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW #surligner les phrases du résumé
    
    create_save_file(doc_result, title=titre, comm=comm, type_file='.docx', move_original=True)
        
    return summary

# In[22]:


def highlight_pdf(summary, titre, verbose=False):
    """
    surligne les phrases de summary dans le pdf titre
    
    paramètres
    -----
    summary : liste de str
        liste des phrases à retrouver dans le pdf
    titre : titre du pdf à retrouver et utiliser pour surligner les phrases
    
    returns
    -----
    copy :
        copie de summary à laquelle on a retiré les phrases trouvées. 
        contient donc les phrases non trouvées dans le pdf (par ex phrases avec ligatures)
    """
    directory = "archives/"
    path = directory + titre
    
    copy = summary.copy()
    l = len(summary)

    doc = fitz.open(path)

    for page in doc:
        text_instances = []
        for idx, phrase in enumerate(summary):
            #print(phrase)
            s = page.search_for(phrase)
            if s != []:#si on trouve la phrase dans la page
                text_instances += s
                try:
                    copy.remove(phrase)
                except:
                    if verbose:
                        #print(phrase)
                        continue

        # HIGHLIGHT
        for inst in text_instances:
            highlight = page.add_highlight_annot(inst)
            highlight.update()
            
    if verbose:
        #print(copy) #affiche les phrases non trouvées dans pdf
        print(f'nb de phrases non trouvées dans le document {titre} = {len(copy)} sur {l}\n') #phrases non trouvées:(((
    create_save_file(doc, title=titre, comm='highlight', type_file='.pdf', move_original=True)
    return copy


# In[23]:


def summarize_pdf(text, titre, comm=None, dictionary = None, n_sentences=5, ratio=0, nlp=None, method: Union[Literal["local"], Literal["global"]] = 'local'):
    '''
    créé un résumé + surligne les phrases sur le pdf
    
    paramètres
    -----
    text:string
        un texte du corpus, le texte à résumer
    titre : str
        titre exact du pdf
    comm : str
        commentaire à ajouter au titre
    dictionary:
        dictionnaire sur tout le corpus
    n_sentences:int
        nombre de phrases max souhaitées pour le résumé
    ratio:float entre 0 et 1
        proportion de phrases souhaitée pour le résumé
    method:'global' ou 'local'
        si method = local
        Compare les phrases d'un document en interne (tfidf calculé par rapport à un doc seulement)
        si method = global
        Compare les phrases par rapport à tous les doc du corpus (sur lequel le dictionnaire a été construit)
        
    nlp : str
        utilisation d'une base pré entrainée et PLUS de textrank
    
    returns
    -----
    summary : résumé créé par summarize_tr
    '''
    
    sentences, summary, sentences_id = summarize_tr(text, dictionary, n_sentences, ratio, method, nlp=nlp)

    highlight_pdf(summary, titre, verbose=False)
        
    return summary

# # 7 Keywords <a name="keywords"></a>
# 
# [Retour table des matières](#up)

# In[24]:


def tfidf_words(corpus, no_below=None, no_above=None, keep_n=None):
    
    '''
    ????????
    Pour calculer le tfidf qui servira à calculer les keywords
    
    corpus = liste de string (liste de tous les documents)
    renvoie le tableau tfidf (taille_corpus,vocab)
    '''
    
    tokenizer = Tokenize(stopwords=["fr"],sentence_tokenize=False, word_tokenize=True, strip_accents=False, lower=True, verbose=False)
    tokenizer += Lemmatizer()
    #word_list est la liste des mots par document
    word_list = tokenizer(corpus)
    
    dictionary = gensim.corpora.Dictionary(word_list)
    #print(len(dictionary))
    if (no_below,no_above,keep_n)!=(None,None,None):
        dictionary.filter_extremes(no_below, no_above, keep_n)
        #print(len(dictionary))
    
    
    corpus = [dictionary.doc2bow(sentence) for sentence in word_list]
    tfidf = TfidfModel(corpus) #fit
    X = tfidf[corpus] #transform
    tfidf_sparse = matutils.corpus2csc(X,num_docs=len(X)).T
    
    return tfidf_sparse, dictionary

# In[25]:


#tfidf, dictionary = tfidf_words(corpus)


# In[26]:


#print(tfidf)


# In[27]:


#corpus


# In[28]:


def get_keywords(doc_id, n_keywords, tfidf, dictionary):
    tfidf_array = tfidf[doc_id].toarray()
    ind = tfidf_array[0].argsort()[-n_keywords:][::-1]
    top_words = []
    for i in ind:
        word = dictionary[i]
        top_words.append(word)
    return top_words, ind


# # 8 Word cloud <a name="wc"></a>
# 
# Premier essai, peu développé pour le moment
# 3 méthodes
# 
# - la méthode de base de la librairie wordcloud (basée uniquement sur la frequence ?)
# - choisir les mots avec le plus haut tfidf
# - choisir les mots avec le plus haut tfidf mais en filtrant le dictionnaire:par exemple supprimer les mots qui apparaissent dans - de 3 documents ou + de 90% des documents du corpus
# 
# [Retour table des matières](#up)

# ## Methode auto (code non modifié) - texte 24

# In[29]:


#cloud2 = data[24]
#print(titres[24])


# In[30]:


'''#automatique
exclure_mots = ['d', 'du', 'de', 'la', 'des', 'le', 'et', 'est', 'elle', 'une', 'en', 'que', 'aux', 'qui', 'ces', 'les', 'dans', 'sur', 'l', 'un', 'pour', 'par', 'il', 'ou', 'à', 'ce', 'a', 'sont', 'cas', 'plus', 'leur', 'se', 's', 'vous', 'au', 'c', 'aussi', 'toutes', 'autre', 'comme']
wordcloud = WordCloud(background_color = 'white', stopwords = exclure_mots, max_words = 50).generate(cloud)
plt.figure(figsize=(10,5))
plt.imshow(wordcloud)
plt.axis("off")
plt.show();'''


# ## mots avec le plus grand tfidf

# In[31]:


'''keywords = get_keywords(109, 50, tfidf, dictionary)[0]
print(keywords)
wc = WordCloud(background_color = 'white')
img = wc.generate_from_text(' '.join(keywords))
plt.figure(figsize=(10,5))
plt.imshow(img)
plt.axis("off")
plt.show();'''


# ## mots avec le plus grand tfidf, dictionnaire filtré

# In[32]:


'''keywords_filter = get_keywords(109, 50, tfidf_filter, dictionary_filter)[0]
print(keywords_filter)
wc = WordCloud(background_color = 'white')
img = wc.generate_from_text(' '.join(keywords_filter))
plt.figure(figsize=(10,5))
plt.imshow(img)
plt.axis("off")
plt.show();

wc.to_file('wordcloud.png')'''


# In[33]:


'''from wordcloud import WordCloud
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
document = Document()

text = data[24]
exclure_mots = ['d', 'du', 'de', 'la', 'des', 'le', 'et', 'est', 'elle', 'une', 'en', 'que', 'aux', 'qui', 'ces', 'les', 'dans', 'sur', 'l', 'un', 'pour', 'par', 'il', 'ou', 'à', 'ce', 'a', 'sont', 'cas', 'plus', 'leur', 'se', 's', 'vous', 'au', 'c', 'aussi', 'toutes', 'autre', 'comme']
wordcloud = WordCloud(background_color = 'white', stopwords = exclure_mots, max_words = 50).generate(data[120])

wordcloud.to_file('wordcloud.png')'''


# # Bert Keywords <a name="bertk"></a>
# 
# [Retour table des matières](#up)
# 
# https://towardsdatascience.com/keyword-extraction-with-bert-724efca412ea

# In[34]:


def embed_kw(doc, nlp, top_n = 20, n_gram_range = (1, 1), stopw=True):
    """
    sert à obtenir les keywords d'un document un utilisant un embedding.
    retourne les mots dont l'embedding est le plus proche de celui du document (cosine similarity)
    
    autre methode?
    
    paramètres 
    -----
    doc : str 
        document dont on veut extraire les keywords
    nlp :
        modèle pré entrainé
    top_n : int
        nb de mots à extraire
    n_gram_range: tuple
        taille des keywords
    stopw : bool
        utilisation des stopwords
        
    returns
    -----
    keywords : list of str
        liste des mots clés
    """
    if stopw:
        stop_words = stopwords.words('french')
    else:
        stop_words = None
    
    #création du bag of words
    count = CountVectorizer(ngram_range=n_gram_range, stop_words=stop_words).fit([doc])
    #candidates = liste des mots
    candidates = count.get_feature_names_out()
    model = nlp
    doc_embedding = model.encode([doc])
    candidate_embeddings = model.encode(candidates)
    distances = cosine_similarity(doc_embedding, candidate_embeddings)
    #on trie les mots selon leur distance au document
    idx = [index for index in distances.argsort()[0][-top_n:]]
    keywords = [candidates[index] for index in distances.argsort()[0][-top_n:]]
    return keywords[::-1]


# In[35]:


#embed_kw(data[0], nlp, top_n = 20, n_gram_range = (1, 1), stopw=True)


# In[36]:


#k = get_keywords(0, 20, tfidf_filter, dictionary_filter)[0]
#print(k)


# In[37]:


def concat_lists(*lists):
    '''
    concatener des listes de mots clés en mettant d'abord le premier mot de chaque liste
    puis le second
    ainsi de suite
    s'arrête à la fin de la liste la plus courte
    
    paramètres
    -----
    *lists : listes à assembler
    
    returns
    ------
    result : liste des mots concaténés
    '''
    Z = zip(*lists)
    result = [x for item in Z for x in item]
    return result


# # 9 Compiled tracts CSV <a name="csv"></a>
# 
# [Retour table des matières](#up)
# 
# [Retour table des matières](#up)

# In[38]:


#compiled_tracts = pd.read_csv('compiled_tracts.csv', '|')


# In[39]:


#compiled_tracts.size


# In[40]:


#compiled_tracts['direction_id'].unique()


# In[41]:


#compiled_tracts.columns


# In[42]:


#compiled_tracts.head(2)


# # 10 import module ML veille sociale <a name="ml"></a>
# 
# [Retour table des matières](#up)

# In[43]:


#!ln -s /home/mathildedacruz/veille-sociale-synthese/machine_learning.py topic/machine_learning.py


# In[44]:


#import module.machine_learning as ml


# In[45]:


#importlib.reload(ml)


# In[46]:


"""x=tracts[0]
sentences = sentence_tok(x)"""


# In[47]:


"""vectorizer = TfidfVectorizer()
X = vectorizer.fit_transform(sentences)
terms = vectorizer.get_feature_names_out()
print(X.shape)"""


# In[48]:


"""nlp = spacy.load("fr_core_news_lg")
nb_top_terms=10"""


# In[49]:


#k, W, H = ml.best_model(X, terms, nlp, nb_top_terms=nb_top_terms, kmin=2, kmax=15)


# In[50]:


#import pygal


# In[51]:


"""radar_theme = ml.top_terms_sorted_pre_pro_graph(H, terms, nb_top_terms)
radar_doc = ml.doc_pre_pro_graph(W)
radar_doc_term = ml.doc_top_terms_sorted_pre_pro_graph(W, H, terms, nb_top_terms)

theme_doc = [np.argmax(doc) for doc in W]
line_chart = pygal.Bar()
line_chart.title = 'nbr de document par Theme'
line_chart.x_labels = ["Theme {}".format(i+1) for i in range(k)]
[line_chart.add("Theme {}".format(i+1), [0] * i+[theme_doc.count(i)]+[0]*(k-i-1)) for i in range(k)]
bar_plot_theme=line_chart.render_data_uri()

#resultatsss = {"states": [{"name": "{} (score: {})".format(titre, round(score, 2)), "command": titre}
 #                         for titre, score
  #                        in zip([titres[i] for i in idx_rearange], requete[IndexMostPertinent])]}

resultats = {"radar_theme": radar_theme,
             "radar_doc": radar_doc,
             "radar_doc_term": radar_doc_term,
             "nb_theme": k,
             "titres": titres}"""


# In[52]:


#resultats['radar_doc']


# # 11 Extract headers <a name="head"></a>
# 
# [Retour table des matières](#up)

# In[53]:


from operator import itemgetter


# In[54]:


def fonts(doc, granularity=False):
    """Extracts fonts and their usage in PDF documents.
   :param doc: PDF document to iterate through
   :type doc: <class 'fitz.fitz.Document'>
   :param granularity: also use 'font', 'flags' and 'color' to discriminate text
   :type granularity: bool
   :rtype: [(font_size, count), (font_size, count}], dict
   :return: most used fonts sorted by count, font style information
    """
    styles = {}
    font_counts = {}

    for page in doc:
        blocks = page.get_text("dict")["blocks"]
        for b in blocks:  # iterate through the text blocks
            if b['type'] == 0:  # block contains text
                for l in b["lines"]:  # iterate through the text lines
                    for s in l["spans"]:  # iterate through the text spans
                        if granularity:
                            identifier = "{0}_{1}_{2}_{3}".format(s['size'], s['flags'], s['font'], s['color'])
                            styles[identifier] = {'size': s['size'], 'flags': s['flags'], 'font': s['font'],
                                                  'color': s['color']}
                        else:
                            identifier = "{0}".format(s['size'])
                            styles[identifier] = {'size': s['size'], 'font': s['font']}

                        font_counts[identifier] = font_counts.get(identifier, 0) + 1  # count the fonts usage

    font_counts = sorted(font_counts.items(), key=itemgetter(1), reverse=True)

    if len(font_counts) < 1:
        raise ValueError("Zero discriminating fonts found!")

    return font_counts, styles


# In[55]:


#path = 'archives/'+titres[1]


# In[56]:


#doc=fitz.open(path)


# In[57]:


#font_counts, styles = fonts(doc, granularity=False)


# In[58]:


#print(font_counts, styles)


# In[59]:


def font_tags(font_counts, styles):
    """Returns dictionary with font sizes as keys and tags as value.
   :param font_counts: (font_size, count) for all fonts occuring in document
   :type font_counts: list
   :param styles: all styles found in the document
   :type styles: dict
   :rtype: dict
   :return: all element tags based on font-sizes
    """
    p_style = styles[font_counts[0][0]]  # get style for most used font by count (paragraph)
    p_size = p_style['size']  # get the paragraph's size

    # sorting the font sizes high to low, so that we can append the right integer to each tag 
    font_sizes = []
    for (font_size, count) in font_counts:
        font_sizes.append(float(font_size))
    font_sizes.sort(reverse=True)

    # aggregating the tags for each font size
    idx = 0
    size_tag = {}
    for size in font_sizes:
        idx += 1
        if size == p_size:
            idx = 0
            size_tag[size] = '<p>'
        if size > p_size:
            size_tag[size] = '<h{0}>'.format(idx)
        elif size < p_size:
            size_tag[size] = '<s{0}>'.format(idx)

    return size_tag


# In[60]:


#size_tag = font_tags(font_counts, styles)


# In[61]:


#size_tag


# In[62]:


def headers_para(doc, size_tag):
    """Scrapes headers & paragraphs from PDF and return texts with element tags.
   :param doc: PDF document to iterate through
   :type doc: <class 'fitz.fitz.Document'>
   :param size_tag: textual element tags for each size
   :type size_tag: dict
   :rtype: list
   :return: texts with pre-prended element tags
    """
    header_para = []  # list with headers and paragraphs
    first = True  # boolean operator for first header
    previous_s = {}  # previous span

    for page in doc:
        blocks = page.getText("dict")["blocks"]
        for b in blocks:  # iterate through the text blocks
            if b['type'] == 0:  # this block contains text

                # REMEMBER: multiple fonts and sizes are possible IN one block

                block_string = ""  # text found in block
                for l in b["lines"]:  # iterate through the text lines
                    for s in l["spans"]:  # iterate through the text spans
                        if s['text'].strip():  # removing whitespaces:
                            if first:
                                previous_s = s
                                first = False
                                block_string = size_tag[s['size']] + s['text']
                            else:
                                if s['size'] == previous_s['size']:

                                    if block_string and all((c == "|") for c in block_string):
                                        # block_string only contains pipes
                                        block_string = size_tag[s['size']] + s['text']
                                    if block_string == "":
                                        # new block has started, so append size tag
                                        block_string = size_tag[s['size']] + s['text']
                                    else:  # in the same block, so concatenate strings
                                        block_string += " " + s['text']

                                else:
                                    header_para.append(block_string)
                                    block_string = size_tag[s['size']] + s['text']

                                previous_s = s

                    # new block started, indicating with a pipe
                    # block_string += "|"

                header_para.append(block_string)

    return header_para


# In[63]:


#text = headers_para(doc, size_tag)


# In[64]:


#text


# In[65]:


#for line in text:
 #   if line.startswith('<h'):
  #      print(line)


# # 12 With embedding from fr_core_news_sm.load() <a name="embed"></a>
# 
# [Retour table des matières](#up)

# In[66]:


def sentence_embedding(sentence, nlp):
    """
    créé un embbedding d'une phrase par la moyenne des embeddings des mots
    
    paramètres
    -----
    sentence : str
        phrase à plonger
    nlp : modèle pré entrainé
    """
    word_embed = nlp(sentence)
    sent_embed = np.mean([word.vector for word in word_embed], axis=0)
    return sent_embed


# In[67]:


def similarity_matrix(doc, nlp):
    """
    Créé la matric de similarité entre toutes les phrases d'un document
    
    doc:list of sentences
    nlp:modele pré entrainé
    
    returns
    -----
    sim : 
    """
    list_embeddings = []
    for sentence in doc:
        list_embeddings.append(sentence_embedding(sentence, nlp))
    embeddings = np.array(list_embeddings)
    sim = cosine_similarity(embeddings,embeddings)
    sim[sim<0]=0
    np.fill_diagonal(sim, 0)
    return sim


# In[68]:


def pdf_highlight(summary,i):
    if titres[i].endswith('pdf'):
        return highlight_pdf(summary,i,verbose=False)


# In[69]:


def get_summaries(data, titres, dictionary, nlp=None, n_sentences=5, ratio=0, highlight = True):
    summaries = []
    for i in tqdm(range(len(data))):
        try:
            summary = summarize_docx(data[i], titre=titres[i], comm='lg', nlp=nlp, dictionary=dictionary, n_sentences=n_sentences, ratio=ratio, method='global')
        except Exception as e:
            summary = []
            print('echec :',e)
        if highlight == True:
            if titres[i].endswith('pdf'):
                highlight_pdf(summary,titres[i],verbose=False)
        summaries.append(summary)
    #print(f'{len(data)-lg} échecs')
    return summaries
    #1, 2, 22, 25, 55, 68, 73, 92, 106, 108, 118, 124


# In[70]:


def get_infos(titres):
    syndicats=[]
    cellules=[]
    for t in titres:
        split = t.split("_")
        syndicats.append(split[0])
        cellules.append(split[1])
    return syndicats, cellules


# # 13 CamemBERT <a name="camembert"></a>
# 
# [Retour table des matières](#up)

# In[71]:


'''from summarizer import Summarizer
from transformers import *

# You can replace "camembert-base" with any other model from the table, e.g. "camembert/camembert-large".
custom_config = AutoConfig.from_pretrained('camembert-base')
custom_config.output_hidden_states=True
tokenizer = CamembertTokenizer.from_pretrained("camembert-base")
camembert = CamembertModel.from_pretrained("camembert-base", config=custom_config)
model = Summarizer(custom_model=camembert, custom_tokenizer=tokenizer)
result = model(tracts[0], num_sentences=5)
full = "".join(result)
print(full)'''


# # 14 keywords embed <a name="kw_embed"></a>
# 
# [Retour table des matières](#up)

# In[72]:


#df_tracts.head(5)


# In[73]:


def list_keywords(data, nlp, n_keywords, tfidf, dictionary):
    """
    retourne des mots clés issu d'une recherche par embedding, et tfidf
    """
    n_keywords = int(n_keywords/2)
    keywords=[]
    for i in tqdm(range(len(data))):
        k1 = get_keywords(i, n_keywords, tfidf, dictionary)[0]
        k2 = embed_kw(data[i], nlp=nlp, top_n = n_keywords, n_gram_range = (1, 1), stopw=True)
        k = concat_lists(k1,k2)
        for word in k:
            if len(word)<3:
                k.remove(word)
        keywords.append(k)
    return keywords


# # 15 Fiche <a name="fiche"></a>
# 
# [Retour table des matières](#up)

# In[74]:


def fiche(df, list_id, titre='corpus', display_wordcloud = True):
    '''
    Utilise le dataframe qui contient déjà toutes les informations pour construire une fiche résumant
    les tracts des id donnés
    
    Affiche les résumés de chaque doc en surlignant les mots clés
    '''
    document = Document()
    concatenate = []
    n = len(list_id)
    for i in list_id:
        keywords = df.iloc[i]['keywords']
        l = len(keywords)//n
        concatenate.append(keywords[:l])
        document.add_heading(df.iloc[i]['titre'])
        summary = df.iloc[i]['resume']
        for phrase in summary:
            paragraph = document.add_paragraph()
            p = re.split(r'(\s)', phrase)
            for word in p:
                run = paragraph.add_run(word)
                if word in keywords:
                    run.bold = True
                    
        if display_wordcloud :
            wc = WordCloud(background_color = 'white')
            img = wc.generate_from_text(' '.join(keywords))
            name = 'wordcloud.png'
            wc.to_file(name)
            run = paragraph.add_run()
            run.add_picture(name, width=Inches(5))
    
    #Wordcloud général
    document.add_heading('WordCloud corpus')
    wc = WordCloud(background_color = 'white')
    res = concat_lists(*concatenate)
    img = wc.generate_from_text(' '.join(res))
    wc.to_file(name)
    document.add_picture(name)
    create_save_file(document, title=titre, type_file='docx', move_original=False)


#words = document.xpath('//w:r', namespaces=document.nsmap)


#words = document.xpath('//w:r', namespaces=document.nsmap)


# In[75]:


#fiche(df_tracts, [0], titre='corpus')


# In[76]:


#fiche(df_tracts, [109,126,32,6,1,110,77], titre='corpus')


# In[77]:


#titres[77]


# # A FAIRE
# 
# - [x] créer fonction filtre
# - [x] keywords avec Bert
# - [ ] ajouter headers
# - [ ] améliorer titres (syndicat, date etc)
# - [ ] metadate
# 
# 
# - [ ] trouver exemple sympa de filtre à montrer

# # 16 Filtres <a name="filtres"></a>
# 
# [Retour table des matières](#up)

# In[78]:


def filter_id(df, list_syndic=None, list_cell=None, start_date=None,end_date=None):
    '''
    Remplir
    '''
    
    if (start_date) and (end_date):
        mask = (df['date'] > start_date) & (df['date'] <= end_date)
        df = df.loc[mask]
    if not list_syndic: #si aucun syndicat n'est renseigné on les prend tous
        list_syndic = df['syndicat'].unique() 
    if not list_cell:
        list_cell = df['cellule'].unique()     
    idx = []
    for i in df.index.values:
        if (df.loc[i]['syndicat'] in list_syndic) and (df.loc[i]['cellule'] in list_cell) :
            idx.append(i)
    if not idx:
        print("Aucun tract correspondant")
    return(idx)




def build_dataframe(directory, n_sentences, n_keywords, ratio=0):
    titres, raw_text, datetime, type_doc = import_data(directory, verbose = False)
    #Appliquer le préprocessing à tous les textes
    data = list(map(text_preproc, raw_text))

    #Création d'un dataframe pandas qui regroupe toutes les info qu'on vient d'extraire
    df_tracts = pd.DataFrame(
        {'titre': titres,
         'texte': data,
         'date' : datetime,
         'type' : type_doc
        })
    sentence_tok, word_tok = init_tokenizer()
    
    dictionary = build_dict_sentence(data, sentence_tok, word_tok)
    #doesn't keep words that appear in less than no_below sentences or more than no_above (ratio) sentences
    dictionary.filter_extremes(no_below=1, no_above=0.1, keep_n=100000)
    
    #tfidf, dictionary = tfidf_words(data)
    tfidf_filter, dictionary_filter = tfidf_words(data, no_below=2, no_above=1)
    
    nlp = SentenceTransformer('paraphrase-multilingual-MiniLM-L12-v2')
    nlp_lg = fr_core_news_lg.load()
    
    print('Calcul des résumés')
    summaries = get_summaries(data, titres, dictionary, nlp=nlp_lg, n_sentences=n_sentences, ratio=ratio)
    
    df_tracts['resume'] = summaries
    
    syndicats, cellules = get_infos(titres)
    df_tracts['syndicat'] = syndicats
    df_tracts['cellule'] = cellules
   
    print('Calcul des mots clés')
    keywords = list_keywords(data, nlp, n_keywords, tfidf_filter, dictionary_filter)
    
    df_tracts['keywords'] = keywords
    
    return df_tracts

